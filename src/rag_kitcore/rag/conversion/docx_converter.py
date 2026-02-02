from docx import Document

from rag_kitcore.logsys.logger import get_logger
from rag_kitcore.core.exceptions import DocumentConversionError
from rag_kitcore.rag.conversion.base import BaseConverter


logger = get_logger(__name__)

class DocxConverter(BaseConverter):

    def convert(self, path: str) -> str:
        return _convert_docx_to_markdown(path)
    
    def __call__(self, path: str) -> str:
        return self.convert(path)


def _convert_docx_to_markdown(file_path: str) -> str:
    """
    Convert a DOCX file to markdown using python-docx.
    Handles headings, paragraphs, bold/italic, and tables.
    Raises DocumentConversionError on failure.
    """
    logger.info("Starting DOCX to Markdown conversion for %s", file_path)

    try:
        doc = Document(file_path)
        md_lines = []

        # --- Paragraphs ---
        for p in doc.paragraphs:
            text = _convert_paragraph_to_markdown(p)
            if text.strip():
                md_lines.append(text)

        # --- Tables ---
        for table in doc.tables:
            md_lines.extend(_convert_table_to_markdown(table))

        markdown_text = "\n".join(md_lines)
        logger.info(
            "DOCX to Markdown conversion completed. Raw length: %d chars",
            len(markdown_text),
        )
        return markdown_text

    except Exception as exc:
        logger.exception("Error during DOCX to Markdown conversion")
        raise DocumentConversionError(
            f"Failed to convert DOCX '{file_path}' to markdown"
        ) from exc


# ---------------------------------------------------------
# Helpers
# ---------------------------------------------------------

def _convert_paragraph_to_markdown(paragraph):
    """Convert a python-docx paragraph into Markdown."""
    style = paragraph.style.name.lower()
    text = _convert_runs_to_markdown(paragraph.runs)

    # Headings
    if "heading" in style:
        try:
            level = int(style.replace("heading", "").strip())
            return "#" * level + " " + text
        except ValueError:
            pass  # fallback to normal paragraph

    # Lists
    if paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.numPr
        if numPr is not None:
            # Numbered list
            return f"1. {text}"

    if paragraph.style.name.lower().startswith("list"):
        # Bullet list
        return f"- {text}"

    return text


def _convert_runs_to_markdown(runs):
    """Convert bold/italic runs to Markdown."""
    parts = []
    for run in runs:
        t = run.text
        if not t:
            continue

        if run.bold:
            t = f"**{t}**"
        if run.italic:
            t = f"*{t}*"

        parts.append(t)

    return "".join(parts)


def _convert_table_to_markdown(table):
    """Convert a python-docx table into Markdown table syntax."""
    md = []

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        md.append("| " + " | ".join(cells) + " |")

        # Header separator after first row
        if i == 0:
            md.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return md

